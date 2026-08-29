# -*- coding: utf-8 -*-
"""Markdown → docx 转换器：标题/表格/列表/粗体/引用/段落。
用法：python md2docx.py <输入.md> <输出.docx> <文档标题>
依赖：python-docx（本机已装，站点构建使用中）。"""
import io
import re
import sys

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def add_runs(par, text):
    """处理 **粗体** 与 `代码` 的行内格式。"""
    pos = 0
    for mo in re.finditer(r'\*\*(.+?)\*\*|`([^`]+)`', text):
        if mo.start() > pos:
            par.add_run(text[pos:mo.start()])
        if mo.group(1) is not None:
            r = par.add_run(mo.group(1))
            r.bold = True
        else:
            r = par.add_run(mo.group(2))
            r.font.name = 'Consolas'
        pos = mo.end()
    if pos < len(text):
        par.add_run(text[pos:])


def convert(md_path, docx_path, title):
    lines = io.open(md_path, encoding='utf-8').read().split('\n')
    doc = Document()
    doc.add_heading(title, level=0)

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        if not stripped:
            i += 1
            continue
        # 表格块
        if stripped.startswith('|') and i + 1 < len(lines) and re.match(r'^\|[\s\-|:]+\|$', lines[i + 1].strip()):
            header = [c.strip() for c in stripped.strip('|').split('|')]
            rows = []
            j = i + 2
            while j < len(lines) and lines[j].strip().startswith('|'):
                rows.append([c.strip() for c in lines[j].strip().strip('|').split('|')])
                j += 1
            table = doc.add_table(rows=1, cols=len(header))
            table.style = 'Light Grid Accent 1'
            for k, h in enumerate(header):
                cell = table.rows[0].cells[min(k, len(header) - 1)]
                cell.text = ''
                add_runs(cell.paragraphs[0], h)
            for row in rows:
                cells = table.add_row().cells
                for k, val in enumerate(row):
                    if k < len(cells):
                        cells[k].text = ''
                        add_runs(cells[k].paragraphs[0], val)
            i = j
            continue
        # 标题
        mo = re.match(r'^(#{1,4})\s+(.*)', stripped)
        if mo:
            level = min(len(mo.group(1)), 3)
            h = doc.add_heading('', level=level)
            add_runs(h, mo.group(2))
            i += 1
            continue
        # 引用
        if stripped.startswith('>'):
            par = doc.add_paragraph(style='Intense Quote')
            add_runs(par, stripped.lstrip('> ').strip())
            i += 1
            continue
        # 无序列表
        mo = re.match(r'^[-*]\s+(.*)', stripped)
        if mo:
            par = doc.add_paragraph(style='List Bullet')
            add_runs(par, mo.group(1))
            i += 1
            continue
        # 有序列表
        mo = re.match(r'^(\d+)\.\s+(.*)', stripped)
        if mo:
            par = doc.add_paragraph(style='List Number')
            add_runs(par, mo.group(2))
            i += 1
            continue
        # 普通段落
        par = doc.add_paragraph()
        add_runs(par, stripped)
        i += 1

    doc.save(docx_path)
    print('DOCX:', docx_path)


if __name__ == '__main__':
    convert(sys.argv[1], sys.argv[2], sys.argv[3])
